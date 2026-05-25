# -*- coding: utf-8 -*-
"""
File Processor Service
پردازش فایل‌های آپلود شده
"""

import os
import io
import uuid
import logging
from datetime import datetime
from typing import Dict, Any, List, Optional, BinaryIO
from pathlib import Path
import pandas as pd

# PDF Processing
try:
    import pdfplumber
    PDF_AVAILABLE = True
except:
    PDF_AVAILABLE = False

# Word Processing
try:
    from docx import Document as DocxDocument
    WORD_AVAILABLE = True
except:
    WORD_AVAILABLE = False

from services.persian_embedding_service import PersianEmbeddingClient, HeydariEmbeddingClient
from processors.advanced_pdf_table_processor import AdvancedPDFTableProcessor

logger = logging.getLogger(__name__)


class FileProcessor:
    """پردازش فایل‌های مختلف"""
    
    def __init__(
        self,
        upload_dir: str = "/home/user01/qwen-api/enhanced_rag_system_dev/uploads",
        collection_manager = None
    ):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True, parents=True)
        
        self.collection_manager = collection_manager
        
        # Initialize embedding service (heydariAI/persian-embeddings)
        try:
            self.embedding_service = HeydariEmbeddingClient()
        except Exception as e:
            logger.warning(f"Failed to initialize HeydariEmbeddingClient: {e}")
            try:
                self.embedding_service = PersianEmbeddingClient()
            except Exception as e2:
                logger.warning(f"Failed to initialize any embedding service: {e2}")
                self.embedding_service = None
        
        # Initialize PDF processor
        self.pdf_processor = AdvancedPDFTableProcessor()
        
        logger.info("✅ FileProcessor initialized")
    
    def save_uploaded_file(
        self,
        file_content: bytes,
        filename: str,
        collection_name: str
    ) -> Dict[str, Any]:
        """ذخیره فایل آپلود شده"""
        try:
            # Generate unique file ID
            file_id = f"{collection_name}_{uuid.uuid4().hex[:12]}"
            
            # Create collection upload dir
            collection_dir = self.upload_dir / collection_name
            collection_dir.mkdir(exist_ok=True, parents=True)
            
            # Save file
            file_path = collection_dir / f"{file_id}_{filename}"
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            logger.info(f"✅ File saved: {file_path}")
            
            return {
                "success": True,
                "file_id": file_id,
                "filename": filename,
                "file_path": str(file_path),
                "file_size": len(file_content)
            }
            
        except Exception as e:
            logger.error(f"❌ Error saving file: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def process_pdf(
        self,
        file_path: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        extract_tables: bool = True
    ) -> Dict[str, Any]:
        """پردازش فایل PDF"""
        try:
            if not PDF_AVAILABLE:
                return {
                    "success": False,
                    "error": "PDF processing not available (pdfplumber not installed)"
                }

            # Read PDF bytes once
            with open(file_path, 'rb') as f:
                pdf_bytes = f.read()

            # ── 1. Extract plain text via pdfplumber ──────────────────────
            full_text_parts = []
            page_count = 0
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        full_text_parts.append(page_text)

            text = "\n\n".join(full_text_parts)

            # ── 2. Extract tables via AdvancedPDFTableProcessor ───────────
            tables = []
            if extract_tables:
                try:
                    tables = self.pdf_processor.extract_tables_advanced(pdf_bytes)
                except Exception as tbl_err:
                    logger.warning(f"Table extraction failed (non-fatal): {tbl_err}")

            # ── 3. Create text chunks ─────────────────────────────────────
            chunks = self._create_chunks(
                text,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap
            )

            # ── 4. Build metadata list ────────────────────────────────────
            metadata_list = [
                {
                    "source": "pdf",
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "page_count": page_count,
                    "has_tables": len(tables) > 0,
                    "tables_count": len(tables),
                    "filename": os.path.basename(file_path),
                }
                for i in range(len(chunks))
            ]

            logger.info(f"✅ PDF processed: {len(chunks)} chunks, {len(tables)} tables, {page_count} pages")

            return {
                "success": True,
                "documents": chunks,
                "metadata_list": metadata_list,
                "tables": tables,
                "text_length": len(text),
                "page_count": page_count,
            }

        except Exception as e:
            logger.error(f"❌ Error processing PDF: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def process_excel(
        self,
        file_path: str,
        chunk_size: int = 500
    ) -> Dict[str, Any]:
        """پردازش فایل Excel"""
        try:
            # Read Excel
            df = pd.read_excel(file_path)
            
            # Convert to text
            documents = []
            metadata_list = []
            
            # Process by rows
            for idx, row in df.iterrows():
                # Create text from row
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                
                if row_text:
                    documents.append(row_text)
                    metadata_list.append({
                        "source": "excel",
                        "row_index": idx,
                        "columns": list(df.columns)
                    })
            
            # Alternative: Convert entire sheet to text
            if len(documents) > 100:  # اگر ردیف‌ها زیاد بود
                # Create summary chunks
                summary_text = df.to_string()
                chunks = self._create_chunks(summary_text, chunk_size=chunk_size)
                
                documents = chunks
                metadata_list = [
                    {"source": "excel", "chunk_index": i, "total_chunks": len(chunks)}
                    for i in range(len(chunks))
                ]
            
            logger.info(f"✅ Excel processed: {len(documents)} documents")
            
            return {
                "success": True,
                "documents": documents,
                "metadata_list": metadata_list,
                "rows_count": len(df),
                "columns_count": len(df.columns)
            }
            
        except Exception as e:
            logger.error(f"❌ Error processing Excel: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def process_csv(
        self,
        file_path: str,
        chunk_size: int = 500
    ) -> Dict[str, Any]:
        """پردازش فایل CSV"""
        try:
            # Read CSV
            df = pd.read_csv(file_path)
            
            # Use same logic as Excel
            return self.process_excel(file_path, chunk_size)
            
        except Exception as e:
            logger.error(f"❌ Error processing CSV: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def process_text(
        self,
        file_path: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ) -> Dict[str, Any]:
        """پردازش فایل متنی"""
        try:
            # Read text file
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # Create chunks
            chunks = self._create_chunks(
                text,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap
            )
            
            # Create metadata
            metadata_list = [
                {"source": "text", "chunk_index": i, "total_chunks": len(chunks)}
                for i in range(len(chunks))
            ]
            
            logger.info(f"✅ Text processed: {len(chunks)} chunks")
            
            return {
                "success": True,
                "documents": chunks,
                "metadata_list": metadata_list,
                "text_length": len(text)
            }
            
        except Exception as e:
            logger.error(f"❌ Error processing text: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def process_word(
        self,
        file_path: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ) -> Dict[str, Any]:
        """پردازش فایل Word"""
        try:
            if not WORD_AVAILABLE:
                return {
                    "success": False,
                    "error": "Word processing not available"
                }
            
            # Read Word document
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            text = "\n\n".join([para.text for para in doc.paragraphs if para.text])
            
            # Create chunks
            chunks = self._create_chunks(
                text,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap
            )
            
            # Create metadata
            metadata_list = [
                {"source": "word", "chunk_index": i, "total_chunks": len(chunks)}
                for i in range(len(chunks))
            ]
            
            logger.info(f"✅ Word processed: {len(chunks)} chunks")
            
            return {
                "success": True,
                "documents": chunks,
                "metadata_list": metadata_list,
                "paragraphs_count": len(doc.paragraphs),
                "text_length": len(text)
            }
            
        except Exception as e:
            logger.error(f"❌ Error processing Word: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def process_json(
        self,
        file_path: str
    ) -> Dict[str, Any]:
        """پردازش فایل JSON"""
        try:
            import json
            
            # Read JSON
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert to documents
            documents = []
            metadata_list = []
            
            if isinstance(data, list):
                # List of objects
                for idx, item in enumerate(data):
                    doc_text = json.dumps(item, ensure_ascii=False, indent=2)
                    documents.append(doc_text)
                    metadata_list.append({
                        "source": "json",
                        "item_index": idx,
                        "total_items": len(data)
                    })
            
            elif isinstance(data, dict):
                # Single object or nested structure
                doc_text = json.dumps(data, ensure_ascii=False, indent=2)
                documents = [doc_text]
                metadata_list = [{"source": "json"}]
            
            logger.info(f"✅ JSON processed: {len(documents)} documents")
            
            return {
                "success": True,
                "documents": documents,
                "metadata_list": metadata_list
            }
            
        except Exception as e:
            logger.error(f"❌ Error processing JSON: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _create_chunks(
        self,
        text: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ) -> List[str]:
        """ساخت chunks از متن"""
        if not text:
            return []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + chunk_size
            
            # Find a good breaking point (space, newline, etc.)
            if end < text_length:
                # Look for the last space or newline before the end
                for break_char in ['\n\n', '\n', '. ', '! ', '? ', ' ']:
                    last_break = text.rfind(break_char, start, end)
                    if last_break != -1:
                        end = last_break + len(break_char)
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - chunk_overlap if chunk_overlap > 0 else end
        
        return chunks
    
    async def process_file(
        self,
        file_id: str,
        collection_name: str,
        file_type: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        extract_tables: bool = True
    ) -> Dict[str, Any]:
        """پردازش فایل بر اساس نوع"""
        try:
            # Find file
            collection_dir = self.upload_dir / collection_name
            file_path = None
            
            for f in collection_dir.glob(f"{file_id}_*"):
                file_path = f
                break
            
            if not file_path or not file_path.exists():
                return {
                    "success": False,
                    "error": f"File not found: {file_id}"
                }
            
            # Process based on file type
            if file_type == "pdf":
                result = self.process_pdf(
                    str(file_path),
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap,
                    extract_tables=extract_tables
                )
            elif file_type == "excel":
                result = self.process_excel(str(file_path), chunk_size=chunk_size)
            elif file_type == "csv":
                result = self.process_csv(str(file_path), chunk_size=chunk_size)
            elif file_type == "text":
                result = self.process_text(
                    str(file_path),
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap
                )
            elif file_type == "word":
                result = self.process_word(
                    str(file_path),
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap
                )
            elif file_type == "json":
                result = self.process_json(str(file_path))
            else:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {file_type}"
                }
            
            if not result["success"]:
                return result
            
            # Add documents to collection (async)
            if self.collection_manager:
                add_result = await self.collection_manager.add_documents(
                    collection_name=collection_name,
                    documents=result["documents"],
                    metadata_list=result["metadata_list"]
                )
                
                if not add_result["success"]:
                    return add_result
            
            logger.info(f"✅ File processed and added to collection: {file_id}")
            
            return {
                "success": True,
                "file_id": file_id,
                "collection_name": collection_name,
                "documents_created": len(result["documents"]),
                "embeddings_created": len(result["documents"]),
                "status": "completed",
                "message": "File processed and added to collection successfully"
            }
            
        except Exception as e:
            logger.error(f"❌ Error processing file: {e}")
            return {
                "success": False,
                "error": str(e)
            }
