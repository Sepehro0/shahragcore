# -*- coding: utf-8 -*-
"""
Universal File Processor
پردازشگر جامع برای تمام فرمت‌های فایل

فرمت‌های پشتیبانی‌شده:
  - PDF  → SmartPersianPDFProcessor (متنی + تصویری + OCR)
  - DOCX → python-docx
  - TXT  → UTF-8 / UTF-16 / Latin-1 auto-detect
  - XLSX / XLS → pandas (هر sheet جداگانه پردازش می‌شود)

خروجی:
  تمام فرمت‌ها به chunks تبدیل شده، embed می‌شوند و در ChromaDB ذخیره می‌گردند.
"""

from __future__ import annotations

import io
import logging
import re
import time
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

CHROMA_DB_PATH = "/home/user01/qwen-api/enhanced_rag_system_dev/chroma_db"

# ──────────────────────────────────────────────────────────────────────────────
# Supported extensions
# ──────────────────────────────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".xlsx", ".xls", ".md"}


def get_file_extension(filename: str) -> str:
    return ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""


def is_supported(filename: str) -> bool:
    return get_file_extension(filename) in SUPPORTED_EXTENSIONS


# ──────────────────────────────────────────────────────────────────────────────
# Data structures
# ──────────────────────────────────────────────────────────────────────────────

@dataclass
class ProcessedChunk:
    text: str
    chunk_index: int
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class FileProcessingStats:
    filename: str
    file_type: str
    chunks: int = 0
    text_chunks: int = 0
    table_chunks: int = 0
    ocr_chunks: int = 0
    pages: int = 0
    time: float = 0.0
    errors: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "filename": self.filename,
            "file_type": self.file_type,
            "chunks": self.chunks,
            "text_chunks": self.text_chunks,
            "table_chunks": self.table_chunks,
            "ocr_chunks": self.ocr_chunks,
            "pages": self.pages,
            "time": round(self.time, 2),
            "errors": self.errors,
        }


# ──────────────────────────────────────────────────────────────────────────────
# Text chunker (sentence-aware, shared with web_content_processor)
# ──────────────────────────────────────────────────────────────────────────────

class _SentenceChunker:
    def __init__(self, chunk_size: int = 700, chunk_overlap: int = 100) -> None:
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk(self, text: str) -> List[str]:
        if not text or len(text.strip()) < 30:
            return []
        sentences: List[str] = []
        # ابتدا بر اساس پاراگراف (دو خط خالی یا بیشتر) تقسیم کن
        paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
        for para in paragraphs:
            # خطوط منفرد داخل پاراگراف را جدا کن (فارسی اغلب هر جمله در یک خط)
            lines = [l.strip() for l in para.splitlines() if l.strip()]
            for line in lines:
                # هر خط را بر اساس نقطه پایان جمله تقسیم کن
                parts = re.split(r"(?<=[.!?؟])\s+", line)
                sentences.extend([p.strip() for p in parts if p.strip()])

        chunks: List[str] = []
        current: List[str] = []
        current_len = 0
        for sent in sentences:
            slen = len(sent)
            if current_len + slen > self.chunk_size and current:
                chunk_text = " ".join(current)
                if len(chunk_text.strip()) >= 30:
                    chunks.append(chunk_text)
                # overlap
                overlap_sents: List[str] = []
                overlap_len = 0
                for s in reversed(current):
                    if overlap_len + len(s) < self.chunk_overlap:
                        overlap_sents.insert(0, s)
                        overlap_len += len(s)
                    else:
                        break
                current = overlap_sents
                current_len = sum(len(s) for s in current)
            current.append(sent)
            current_len += slen
        if current:
            chunk_text = " ".join(current)
            if len(chunk_text.strip()) >= 30:
                chunks.append(chunk_text)
        return chunks


# ──────────────────────────────────────────────────────────────────────────────
# DOCX processor
# ──────────────────────────────────────────────────────────────────────────────

def _process_docx(
    file_bytes: bytes,
    filename: str,
    chunk_size: int,
    chunk_overlap: int,
    base_metadata: Dict[str, Any],
    chunk_offset: int,
) -> Tuple[List[ProcessedChunk], FileProcessingStats]:
    stats = FileProcessingStats(filename=filename, file_type="docx")
    t0 = time.time()
    chunks: List[ProcessedChunk] = []

    try:
        from docx import Document as DocxDocument
    except ImportError:
        stats.errors.append("python-docx is not installed")
        stats.time = time.time() - t0
        return chunks, stats

    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        full_text_parts: List[str] = []

        # متن پاراگراف‌ها
        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt:
                full_text_parts.append(txt)

        # جداول
        table_texts: List[str] = []
        for table_idx, table in enumerate(doc.tables):
            rows_text: List[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows_text.append(" | ".join(cells))
            if rows_text:
                table_str = f"[جدول {table_idx + 1}]\n" + "\n".join(rows_text)
                table_texts.append(table_str)

        full_text = "\n\n".join(full_text_parts)
        stats.pages = 1  # DOCX بدون شماره صفحه

        chunker = _SentenceChunker(chunk_size, chunk_overlap)
        raw_chunks = chunker.chunk(full_text)

        for i, chunk_text in enumerate(raw_chunks):
            meta = {
                **base_metadata,
                "chunk_index": chunk_offset + i,
                "page_chunk_index": i,
                "total_page_chunks": len(raw_chunks),
                "char_count": len(chunk_text),
                "content_type": "text",
            }
            chunks.append(ProcessedChunk(text=chunk_text, chunk_index=chunk_offset + i, metadata=meta))

        # اضافه کردن table chunks
        for i, table_text in enumerate(table_texts):
            idx = chunk_offset + len(raw_chunks) + i
            meta = {
                **base_metadata,
                "chunk_index": idx,
                "content_type": "table",
                "char_count": len(table_text),
            }
            chunks.append(ProcessedChunk(text=table_text, chunk_index=idx, metadata=meta))

        stats.chunks = len(chunks)
        stats.text_chunks = len(raw_chunks)
        stats.table_chunks = len(table_texts)

    except Exception as e:
        logger.error(f"DOCX processing failed for {filename}: {e}", exc_info=True)
        stats.errors.append(str(e))

    stats.time = time.time() - t0
    return chunks, stats


# ──────────────────────────────────────────────────────────────────────────────
# TXT processor
# ──────────────────────────────────────────────────────────────────────────────

def _process_txt(
    file_bytes: bytes,
    filename: str,
    chunk_size: int,
    chunk_overlap: int,
    base_metadata: Dict[str, Any],
    chunk_offset: int,
) -> Tuple[List[ProcessedChunk], FileProcessingStats]:
    stats = FileProcessingStats(filename=filename, file_type="txt")
    t0 = time.time()
    chunks: List[ProcessedChunk] = []

    # تشخیص encoding
    content = ""
    for enc in ("utf-8-sig", "utf-8", "utf-16", "windows-1256", "latin-1"):
        try:
            content = file_bytes.decode(enc)
            break
        except (UnicodeDecodeError, ValueError):
            continue

    if not content:
        stats.errors.append("Could not decode file — unknown encoding")
        stats.time = time.time() - t0
        return chunks, stats

    stats.pages = 1
    chunker = _SentenceChunker(chunk_size, chunk_overlap)
    raw_chunks = chunker.chunk(content)

    for i, chunk_text in enumerate(raw_chunks):
        meta = {
            **base_metadata,
            "chunk_index": chunk_offset + i,
            "page_chunk_index": i,
            "total_page_chunks": len(raw_chunks),
            "char_count": len(chunk_text),
            "content_type": "text",
        }
        chunks.append(ProcessedChunk(text=chunk_text, chunk_index=chunk_offset + i, metadata=meta))

    stats.chunks = len(chunks)
    stats.text_chunks = len(chunks)
    stats.time = time.time() - t0
    return chunks, stats


# ──────────────────────────────────────────────────────────────────────────────
# Excel processor
# ──────────────────────────────────────────────────────────────────────────────

def _process_excel(
    file_bytes: bytes,
    filename: str,
    chunk_size: int,
    chunk_overlap: int,
    base_metadata: Dict[str, Any],
    chunk_offset: int,
) -> Tuple[List[ProcessedChunk], FileProcessingStats]:
    stats = FileProcessingStats(filename=filename, file_type="excel")
    t0 = time.time()
    chunks: List[ProcessedChunk] = []

    try:
        import pandas as pd
    except ImportError:
        stats.errors.append("pandas is not installed")
        stats.time = time.time() - t0
        return chunks, stats

    try:
        df_dict = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
        sheet_chunks: List[ProcessedChunk] = []
        idx = chunk_offset

        for sheet_name, df in df_dict.items():
            df = df.fillna("").astype(str)
            # header
            header_row = " | ".join(str(c) for c in df.columns)

            # هر ردیف یک chunk (برای جداول کوچک) یا batch (برای بزرگ)
            rows = df.values.tolist()
            if not rows:
                continue

            # اگر تعداد ردیف‌ها کم باشد: تمام sheet یک chunk
            if len(rows) <= 50:
                rows_text = "\n".join(" | ".join(str(c) for c in row) for row in rows)
                chunk_text = f"[شیت: {sheet_name}]\n{header_row}\n{rows_text}"
                meta = {
                    **base_metadata,
                    "chunk_index": idx,
                    "sheet_name": sheet_name,
                    "content_type": "table",
                    "rows_count": len(rows),
                    "char_count": len(chunk_text),
                }
                sheet_chunks.append(ProcessedChunk(text=chunk_text, chunk_index=idx, metadata=meta))
                idx += 1
            else:
                # batch: هر batch_size ردیف یک chunk
                batch_size = 40
                for b_start in range(0, len(rows), batch_size):
                    batch = rows[b_start : b_start + batch_size]
                    rows_text = "\n".join(" | ".join(str(c) for c in row) for row in batch)
                    chunk_text = (
                        f"[شیت: {sheet_name} - ردیف {b_start + 1} تا {b_start + len(batch)}]\n"
                        f"{header_row}\n{rows_text}"
                    )
                    meta = {
                        **base_metadata,
                        "chunk_index": idx,
                        "sheet_name": sheet_name,
                        "content_type": "table",
                        "rows_count": len(batch),
                        "row_start": b_start + 1,
                        "char_count": len(chunk_text),
                    }
                    sheet_chunks.append(ProcessedChunk(text=chunk_text, chunk_index=idx, metadata=meta))
                    idx += 1

        chunks = sheet_chunks
        stats.chunks = len(chunks)
        stats.table_chunks = len(chunks)
        stats.pages = len(df_dict)

    except Exception as e:
        logger.error(f"Excel processing failed for {filename}: {e}", exc_info=True)
        stats.errors.append(str(e))

    stats.time = time.time() - t0
    return chunks, stats


# ──────────────────────────────────────────────────────────────────────────────
# Main Universal Processor
# ──────────────────────────────────────────────────────────────────────────────

class UniversalFileProcessor:
    """
    پردازشگر جامع فایل‌ها.
    همه فرمت‌ها را به chunks تبدیل کرده و در ChromaDB ذخیره می‌کند.
    """

    _embed_model = None

    def __init__(self, chunk_size: int = 700, chunk_overlap: int = 100) -> None:
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def _get_embed_model(self):
        if UniversalFileProcessor._embed_model is None:
            from services.persian_embedding_service import get_heydari_model
            UniversalFileProcessor._embed_model = get_heydari_model()
        return UniversalFileProcessor._embed_model

    def _get_chroma_client(self):
        import chromadb
        from chromadb.config import Settings as ChromaSettings
        return chromadb.PersistentClient(
            path=CHROMA_DB_PATH,
            settings=ChromaSettings(anonymized_telemetry=False),
        )

    def _embed_chunks(self, chunks: List[ProcessedChunk]) -> List[List[float]]:
        model = self._get_embed_model()
        texts = [c.text for c in chunks]
        batch_size = 32
        embeddings: List[List[float]] = []
        for i in range(0, len(texts), batch_size):
            batch = texts[i : i + batch_size]
            embs = model.encode(batch, show_progress_bar=False).tolist()
            embeddings.extend(embs)
        return embeddings

    def _save_to_chroma(
        self,
        collection_name: str,
        chunks: List[ProcessedChunk],
        embeddings: List[List[float]],
        overwrite: bool,
        append: bool,
    ) -> int:
        client = self._get_chroma_client()

        if overwrite and not append:
            try:
                client.delete_collection(collection_name)
                logger.info(f"  🗑️ Deleted existing collection: {collection_name}")
            except Exception:
                pass

        try:
            collection = client.get_collection(collection_name)
        except Exception:
            try:
                from services.persian_embedding_service import HEYDARI_EMBEDDING_DIM
                emb_dim = HEYDARI_EMBEDDING_DIM
            except Exception:
                emb_dim = len(embeddings[0]) if embeddings else 1024
            collection = client.create_collection(
                name=collection_name,
                metadata={
                    "hnsw:space": "cosine",
                    "embedding_dim": str(emb_dim),
                    "created_at": datetime.utcnow().isoformat(),
                },
            )

        offset = 0
        if append:
            try:
                offset = client.get_collection(collection_name).count()
            except Exception:
                offset = 0

        ids, documents, meta_list, valid_embs = [], [], [], []
        for chunk, emb in zip(chunks, embeddings):
            text = chunk.text.strip()
            if not text or len(text) < 15:
                continue
            ids.append(f"{collection_name}_{offset + chunk.chunk_index}")
            documents.append(text)
            valid_embs.append(emb)
            clean_meta: Dict[str, Any] = {}
            for k, v in chunk.metadata.items():
                if v is None:
                    continue
                clean_meta[k] = v if isinstance(v, (str, int, float, bool)) else str(v)
            meta_list.append(clean_meta)

        batch_size = 100
        saved = 0
        for i in range(0, len(ids), batch_size):
            end = min(i + batch_size, len(ids))
            collection.add(
                ids=ids[i:end],
                documents=documents[i:end],
                embeddings=valid_embs[i:end],
                metadatas=meta_list[i:end],
            )
            saved += end - i
        return saved

    def process_files(
        self,
        files: List[Dict[str, Any]],
        collection_name: str,
        overwrite: bool = True,
        append: bool = False,
        extra_metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        پردازش لیستی از فایل‌ها و ذخیره در ChromaDB.

        Args:
            files: [{"bytes": bytes, "filename": str, "metadata": dict}, ...]
            collection_name: نام collection
            overwrite: پاک کردن collection قبلی
            append: اضافه کردن به collection موجود
            extra_metadata: metadata اضافی برای همه chunks

        Returns:
            {success, total_chunks, total_files, stats_per_file, error}
        """
        t0 = time.time()
        all_chunks: List[ProcessedChunk] = []
        stats_list: List[FileProcessingStats] = []
        chunk_offset = 0
        extra_metadata = extra_metadata or {}

        for file_info in files:
            fb = file_info["bytes"]
            fname = file_info.get("filename", "unnamed")
            file_meta = file_info.get("metadata", {})

            ext = get_file_extension(fname)
            base_meta = {
                "source_file": fname,
                "collection": collection_name,
                "file_type": ext.lstrip("."),
                "source_type": "file_upload",
                "crawl_date": datetime.utcnow().isoformat(),
                "source": fname,
                **extra_metadata,
                **file_meta,
            }

            logger.info(f"  📄 Processing [{ext}]: {fname}")

            if ext == ".pdf":
                # از SmartPersianPDFProcessor استفاده می‌شود
                file_chunks, fstats = self._process_pdf(
                    fb, fname, base_meta, chunk_offset
                )
            elif ext in (".docx", ".doc"):
                file_chunks, fstats = _process_docx(
                    fb, fname, self.chunk_size, self.chunk_overlap, base_meta, chunk_offset
                )
            elif ext == ".txt":
                file_chunks, fstats = _process_txt(
                    fb, fname, self.chunk_size, self.chunk_overlap, base_meta, chunk_offset
                )
            elif ext in (".xlsx", ".xls"):
                file_chunks, fstats = _process_excel(
                    fb, fname, self.chunk_size, self.chunk_overlap, base_meta, chunk_offset
                )
            elif ext == ".md":
                # Markdown → TXT
                file_chunks, fstats = _process_txt(
                    fb, fname, self.chunk_size, self.chunk_overlap, base_meta, chunk_offset
                )
                fstats.file_type = "markdown"
            else:
                logger.warning(f"  ⚠️ Unsupported file type: {fname}")
                fstats = FileProcessingStats(filename=fname, file_type=ext)
                fstats.errors.append(f"Unsupported file type: {ext}")
                stats_list.append(fstats)
                continue

            all_chunks.extend(file_chunks)
            chunk_offset += len(file_chunks)
            stats_list.append(fstats)
            logger.info(
                f"  ✅ {fname}: {fstats.chunks} chunks in {fstats.time:.1f}s"
            )

        if not all_chunks:
            return {
                "success": False,
                "error": "هیچ chunk‌ای از فایل‌ها تولید نشد",
                "total_chunks": 0,
                "total_files": len(files),
                "total_time": round(time.time() - t0, 1),
                "stats_per_file": [s.to_dict() for s in stats_list],
            }

        # Embedding
        logger.info(f"  🔢 Generating embeddings for {len(all_chunks)} chunks...")
        try:
            embeddings = self._embed_chunks(all_chunks)
        except Exception as e:
            return {
                "success": False,
                "error": f"Embedding failed: {e}",
                "total_chunks": 0,
                "total_files": len(files),
                "total_time": round(time.time() - t0, 1),
                "stats_per_file": [s.to_dict() for s in stats_list],
            }

        # ذخیره در ChromaDB
        logger.info(f"  💾 Saving to ChromaDB collection '{collection_name}'...")
        try:
            saved = self._save_to_chroma(
                collection_name, all_chunks, embeddings, overwrite, append
            )
        except Exception as e:
            return {
                "success": False,
                "error": f"ChromaDB save failed: {e}",
                "total_chunks": 0,
                "total_files": len(files),
                "total_time": round(time.time() - t0, 1),
                "stats_per_file": [s.to_dict() for s in stats_list],
            }

        total_time = round(time.time() - t0, 1)
        logger.info(
            f"  🎉 Done: {saved} chunks saved in {total_time}s "
            f"({len(files)} files)"
        )

        return {
            "success": True,
            "total_chunks": saved,
            "total_files": len(files),
            "total_time": total_time,
            "stats_per_file": [s.to_dict() for s in stats_list],
        }

    def _process_pdf(
        self,
        file_bytes: bytes,
        filename: str,
        base_metadata: Dict[str, Any],
        chunk_offset: int,
    ) -> Tuple[List[ProcessedChunk], FileProcessingStats]:
        """پردازش PDF با SmartPersianPDFProcessor."""
        from processors.smart_persian_pdf_processor import SmartPersianPDFProcessor

        stats = FileProcessingStats(filename=filename, file_type="pdf")
        t0 = time.time()
        try:
            processor = SmartPersianPDFProcessor(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
            )
            pdf_chunks, pdf_stats = processor.process_single_pdf(
                pdf_bytes=file_bytes,
                filename=filename,
                collection_name=base_metadata.get("collection", "default"),
                extra_metadata={k: v for k, v in base_metadata.items()
                                if k not in ("source_file", "collection")},
                chunk_offset=chunk_offset,
            )
            # تبدیل ProcessedChunk از smart_persian به UniversalFileProcessor.ProcessedChunk
            universal_chunks = [
                ProcessedChunk(
                    text=c.text,
                    chunk_index=c.chunk_index,
                    metadata=c.metadata,
                )
                for c in pdf_chunks
            ]
            stats.chunks = len(universal_chunks)
            stats.text_chunks = pdf_stats.text_chunks
            stats.table_chunks = pdf_stats.table_chunks
            stats.ocr_chunks = pdf_stats.ocr_chunks
            stats.pages = pdf_stats.total_pages
            stats.errors = pdf_stats.errors
        except Exception as e:
            logger.error(f"PDF processing failed: {e}", exc_info=True)
            stats.errors.append(str(e))
            universal_chunks = []

        stats.time = time.time() - t0
        return universal_chunks, stats
