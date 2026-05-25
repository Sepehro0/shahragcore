# -*- coding: utf-8 -*-
"""
Document Processing Module
پردازشگر اسناد
"""

import os
import io
import logging
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from enum import Enum

# PDF processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PDF processing libraries not available")

# Office documents
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available")

# Excel files
try:
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logging.warning("pandas not available")

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """انواع اسناد"""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    XLSX = "xlsx"
    XLS = "xls"
    MD = "md"
    HTML = "html"
    UNKNOWN = "unknown"


@dataclass
class ProcessedDocument:
    """سند پردازش شده"""
    content: str
    document_type: DocumentType
    metadata: Dict[str, Any]
    pages: List[Dict[str, Any]]
    tables: List[Dict[str, Any]]
    success: bool
    error: Optional[str] = None


class DocumentProcessor:
    """پردازشگر اسناد"""
    
    def __init__(self):
        self.supported_types = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.txt': DocumentType.TXT,
            '.xlsx': DocumentType.XLSX,
            '.xls': DocumentType.XLS,
            '.md': DocumentType.MD,
            '.html': DocumentType.HTML
        }
    
    def detect_document_type(self, filename: str) -> DocumentType:
        """تشخیص نوع سند"""
        if not filename:
            return DocumentType.UNKNOWN
        
        _, ext = os.path.splitext(filename.lower())
        return self.supported_types.get(ext, DocumentType.UNKNOWN)
    
    def process_document(self, file_bytes: bytes, filename: str) -> ProcessedDocument:
        """پردازش سند"""
        try:
            document_type = self.detect_document_type(filename)
            
            if document_type == DocumentType.PDF:
                return self._process_pdf(file_bytes, filename)
            elif document_type == DocumentType.DOCX:
                return self._process_docx(file_bytes, filename)
            elif document_type == DocumentType.TXT:
                return self._process_txt(file_bytes, filename)
            elif document_type in [DocumentType.XLSX, DocumentType.XLS]:
                return self._process_excel(file_bytes, filename)
            elif document_type == DocumentType.MD:
                return self._process_markdown(file_bytes, filename)
            elif document_type == DocumentType.HTML:
                return self._process_html(file_bytes, filename)
            else:
                return ProcessedDocument(
                    content="",
                    document_type=DocumentType.UNKNOWN,
                    metadata={},
                    pages=[],
                    tables=[],
                    success=False,
                    error=f"Unsupported document type: {filename}"
                )
        
        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            return ProcessedDocument(
                content="",
                document_type=DocumentType.UNKNOWN,
                metadata={},
                pages=[],
                tables=[],
                success=False,
                error=str(e)
            )
    
    def _process_pdf(self, file_bytes: bytes, filename: str) -> ProcessedDocument:
        """پردازش PDF"""
        if not PDF_AVAILABLE:
            return ProcessedDocument(
                content="",
                document_type=DocumentType.PDF,
                metadata={},
                pages=[],
                tables=[],
                success=False,
                error="PDF processing libraries not available"
            )
        
        try:
            content = ""
            pages = []
            tables = []
            
            # استفاده از pdfplumber برای استخراج متن و جداول
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # استخراج متن
                    page_text = page.extract_text() or ""
                    content += page_text + "\n"
                    
                    # استخراج جداول
                    page_tables = page.extract_tables()
                    for table_idx, table in enumerate(page_tables):
                        if table:
                            tables.append({
                                'page': page_num + 1,
                                'table_index': table_idx,
                                'data': table
                            })
                    
                    pages.append({
                        'page_number': page_num + 1,
                        'text': page_text,
                        'tables_count': len(page_tables)
                    })
            
            return ProcessedDocument(
                content=content.strip(),
                document_type=DocumentType.PDF,
                metadata={
                    'filename': filename,
                    'total_pages': len(pages),
                    'total_tables': len(tables)
                },
                pages=pages,
                tables=tables,
                success=True
            )
        
        except Exception as e:
            logger.error(f"Error processing PDF {filename}: {e}")
            return ProcessedDocument(
                content="",
                document_type=DocumentType.PDF,
                metadata={},
                pages=[],
                tables=[],
                success=False,
                error=str(e)
            )
    
    def _process_docx(self, file_bytes: bytes, filename: str) -> ProcessedDocument:
        """پردازش DOCX"""
        if not DOCX_AVAILABLE:
            return ProcessedDocument(
                content="",
                document_type=DocumentType.DOCX,
                metadata={},
                pages=[],
                tables=[],
                success=False,
                error="python-docx not available"
            )
        
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            
            content = ""
            tables = []
            
            # استخراج متن
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # استخراج جداول
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                
                tables.append({
                    'table_index': table_idx,
                    'data': table_data
                })
            
            return ProcessedDocument(
                content=content.strip(),
                document_type=DocumentType.DOCX,
                metadata={
                    'filename': filename,
                    'total_tables': len(tables)
                },
                pages=[{
                    'page_number': 1,
                    'text': content,
                    'tables_count': len(tables)
                }],
                tables=tables,
                success=True
            )
        
        except Exception as e:
            logger.error(f"Error processing DOCX {filename}: {e}")
            return ProcessedDocument(
                content="",
                document_type=DocumentType.DOCX,
                metadata={},
                pages=[],
                tables=[],
                success=False,
                error=str(e)
            )
    
    def _process_txt(self, file_bytes: bytes, filename: str) -> ProcessedDocument:
        """پردازش TXT"""
        try:
            content = file_bytes.decode('utf-8')
            
            return ProcessedDocument(
                content=content,
                document_type=DocumentType.TXT,
                metadata={
                    'filename': filename,
                    'encoding': 'utf-8'
                },
                pages=[{
                    'page_number': 1,
                    'text': content,
                    'tables_count': 0
                }],
                tables=[],
                success=True
            )
        
        except UnicodeDecodeError:
            try:
                content = file_bytes.decode('latin-1')
                return ProcessedDocument(
                    content=content,
                    document_type=DocumentType.TXT,
                    metadata={
                        'filename': filename,
                        'encoding': 'latin-1'
                    },
                    pages=[{
                        'page_number': 1,
                        'text': content,
                        'tables_count': 0
                    }],
                    tables=[],
                    success=True
                )
            except Exception as e:
                return ProcessedDocument(
                    content="",
                    document_type=DocumentType.TXT,
                    metadata={},
                    pages=[],
                    tables=[],
                    success=False,
                    error=f"Could not decode text file: {e}"
                )
    
    def _process_excel(self, file_bytes: bytes, filename: str) -> ProcessedDocument:
        """پردازش Excel"""
        if not EXCEL_AVAILABLE:
            return ProcessedDocument(
                content="",
                document_type=DocumentType.XLSX,
                metadata={},
                pages=[],
                tables=[],
                success=False,
                error="pandas not available"
            )
        
        try:
            # خواندن Excel
            df_dict = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
            
            content = ""
            tables = []
            
            for sheet_name, df in df_dict.items():
                # تبدیل DataFrame به متن
                sheet_content = f"Sheet: {sheet_name}\n"
                sheet_content += df.to_string(index=False)
                content += sheet_content + "\n\n"
                
                # تبدیل DataFrame به جدول
                table_data = df.values.tolist()
                table_data.insert(0, df.columns.tolist())  # اضافه کردن header
                
                tables.append({
                    'sheet_name': sheet_name,
                    'data': table_data
                })
            
            return ProcessedDocument(
                content=content.strip(),
                document_type=DocumentType.XLSX,
                metadata={
                    'filename': filename,
                    'total_sheets': len(df_dict),
                    'total_tables': len(tables)
                },
                pages=[{
                    'page_number': 1,
                    'text': content,
                    'tables_count': len(tables)
                }],
                tables=tables,
                success=True
            )
        
        except Exception as e:
            logger.error(f"Error processing Excel {filename}: {e}")
            return ProcessedDocument(
                content="",
                document_type=DocumentType.XLSX,
                metadata={},
                pages=[],
                tables=[],
                success=False,
                error=str(e)
            )
    
    def _process_markdown(self, file_bytes: bytes, filename: str) -> ProcessedDocument:
        """پردازش Markdown"""
        try:
            content = file_bytes.decode('utf-8')
            
            return ProcessedDocument(
                content=content,
                document_type=DocumentType.MD,
                metadata={
                    'filename': filename,
                    'encoding': 'utf-8'
                },
                pages=[{
                    'page_number': 1,
                    'text': content,
                    'tables_count': 0
                }],
                tables=[],
                success=True
            )
        
        except Exception as e:
            logger.error(f"Error processing Markdown {filename}: {e}")
            return ProcessedDocument(
                content="",
                document_type=DocumentType.MD,
                metadata={},
                pages=[],
                tables=[],
                success=False,
                error=str(e)
            )
    
    def _process_html(self, file_bytes: bytes, filename: str) -> ProcessedDocument:
        """پردازش HTML"""
        try:
            content = file_bytes.decode('utf-8')
            
            # استخراج متن از HTML (ساده)
            import re
            # حذف تگ‌های HTML
            text_content = re.sub(r'<[^>]+>', '', content)
            # حذف فضاهای اضافی
            text_content = re.sub(r'\s+', ' ', text_content).strip()
            
            return ProcessedDocument(
                content=text_content,
                document_type=DocumentType.HTML,
                metadata={
                    'filename': filename,
                    'encoding': 'utf-8',
                    'original_html': content
                },
                pages=[{
                    'page_number': 1,
                    'text': text_content,
                    'tables_count': 0
                }],
                tables=[],
                success=True
            )
        
        except Exception as e:
            logger.error(f"Error processing HTML {filename}: {e}")
            return ProcessedDocument(
                content="",
                document_type=DocumentType.HTML,
                metadata={},
                pages=[],
                tables=[],
                success=False,
                error=str(e)
            )


    def extract_text_from_pdf(self, file_bytes: bytes) -> str:
        """استخراج متن از PDF (برای سازگاری)"""
        result = self._process_pdf(file_bytes, "temp.pdf")
        return result.content if result.success else ""


# Global instance
document_processor = DocumentProcessor()
